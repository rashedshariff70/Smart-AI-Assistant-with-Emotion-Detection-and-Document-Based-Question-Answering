from flask import Flask, render_template, request, jsonify, session
from dotenv import load_dotenv
import os
import json
import re
from groq import Groq
from werkzeug.utils import secure_filename
import tempfile

# PDF/DOC reading
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

load_dotenv()

app = Flask(__name__)
app.secret_key = os.urandom(24)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB max upload

# ── Load API key + model from .env ──────────────────────────────────────────
raw_key = os.getenv("GROQ_API_KEY_QWEN", "")
if "#" in raw_key:
    api_key, model_name = raw_key.split("#", 1)
    api_key = api_key.strip()
    model_name = model_name.strip()
else:
    api_key = raw_key.strip()
    model_name = "qwen/qwen3-32b"

client = Groq(api_key=api_key)

UPLOAD_FOLDER = tempfile.mkdtemp()
ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}

doc_store: dict[str, str] = {}
conv_store: dict[str, list] = {}


# ────────────────────────────────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────────────────────────────────

def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text(filepath: str, ext: str) -> str:
    """Extract plain text from PDF / DOCX / TXT."""
    if ext == "txt":
        with open(filepath, "r", errors="ignore") as f:
            return f.read()

    if ext == "pdf":
        if not PDF_SUPPORT:
            raise RuntimeError("PyPDF2 not installed. Run: pip install PyPDF2")
        text = []
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
        return "\n".join(text)

    if ext == "docx":
        if not DOCX_SUPPORT:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
        doc = DocxDocument(filepath)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        # Also grab table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)

    return ""


def clean_llm_output(raw: str) -> str:
    """
    Qwen3 (and some other models) emit <think>...</think> reasoning blocks
    before the actual answer. Strip them plus any markdown fences so we get
    clean text / JSON back.
    """
    # Remove <think>...</think> blocks (including multi-line)
    raw = re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL)
    # Remove ```json or ``` fences
    raw = re.sub(r"```(?:json)?", "", raw)
    raw = raw.strip().strip("`").strip()
    return raw


def call_groq(system_prompt: str, user_message: str, temperature: float = 0.7) -> str:
    """Single-turn Groq call."""
    response = client.chat.completions.create(
        model=model_name,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_message},
        ],
        temperature=temperature,
        max_tokens=1024,
    )
    return response.choices[0].message.content.strip()


# ────────────────────────────────────────────────────────────────────────────
# Routes – Pages
# ────────────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/prediction")
def prediction():
    return render_template("prediction.html")


# ────────────────────────────────────────────────────────────────────────────
# API – Stage 1 : Emotion Detection
# ────────────────────────────────────────────────────────────────────────────

EMOTION_SYSTEM = """You are an expert emotion analysis AI.
Analyze the user's text and respond ONLY with a valid JSON object — no markdown, no code fences, no <think> blocks, no extra explanation.
Return ONLY this JSON and nothing else:
{
  "primary_emotion": "<one of: joy|sadness|anger|fear|surprise|disgust|neutral>",
  "confidence": <integer 0-100>,
  "emotions": {
    "joy": <integer 0-100>,
    "sadness": <integer 0-100>,
    "anger": <integer 0-100>,
    "fear": <integer 0-100>,
    "surprise": <integer 0-100>,
    "disgust": <integer 0-100>,
    "neutral": <integer 0-100>
  },
  "sentiment": "<positive|negative|neutral>",
  "intensity": "<low|medium|high>",
  "summary": "<one sentence empathetic insight about the text>"
}"""


@app.route("/api/emotion", methods=["POST"])
def api_emotion():
    data = request.get_json(silent=True) or {}
    text = (data.get("text") or "").strip()
    if not text:
        return jsonify({"error": "No text provided"}), 400

    try:
        raw = call_groq(EMOTION_SYSTEM, text, temperature=0.3)
        cleaned = clean_llm_output(raw)

        # If model still wrapped in fences, try to extract the JSON object
        if not cleaned.startswith("{"):
            match = re.search(r"\{.*\}", cleaned, re.DOTALL)
            if match:
                cleaned = match.group(0)
            else:
                raise ValueError(f"No JSON object found in model response. Raw: {raw[:300]}")

        result = json.loads(cleaned)
        return jsonify(result)
    except json.JSONDecodeError as e:
        return jsonify({"error": f"JSON parse error: {str(e)}. Try again."}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ────────────────────────────────────────────────────────────────────────────
# API – Stage 2 : RAG Document Chatbot
# ────────────────────────────────────────────────────────────────────────────

@app.route("/api/upload", methods=["POST"])
def api_upload():
    if "file" not in request.files:
        return jsonify({"error": "No file part in request"}), 400

    file = request.files["file"]
    if not file or file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": "Unsupported file type. Please upload a PDF, DOCX, or TXT file."}), 400

    filename = secure_filename(file.filename)
    ext = filename.rsplit(".", 1)[1].lower()
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    try:
        text = extract_text(filepath, ext)
    except RuntimeError as e:
        return jsonify({"error": str(e)}), 500

    if not text.strip():
        return jsonify({"error": "Could not extract any text from this document. The file may be scanned/image-based."}), 400

    sid = session.get("sid") or os.urandom(8).hex()
    session["sid"] = sid
    doc_store[sid] = text[:15000]   # cap to ~15k chars for context window

    return jsonify({"success": True, "chars": len(text), "preview": text[:300]})


@app.route("/api/rag", methods=["POST"])
def api_rag():
    data = request.get_json(silent=True) or {}
    question = (data.get("question") or "").strip()
    sid = session.get("sid", "")

    if not question:
        return jsonify({"error": "No question provided"}), 400

    doc_text = doc_store.get(sid, "")
    if not doc_text:
        return jsonify({"error": "No document uploaded yet. Please upload a PDF, DOCX, or TXT file first."}), 400

    system = f"""You are a precise and helpful document assistant.
Answer the user's question ONLY using information found in the document below.
Be thorough and cite specific parts of the document when relevant.
If the answer is not present in the document, respond with: "I couldn't find that information in the uploaded document."
Do not hallucinate or use outside knowledge.

DOCUMENT CONTENT:
{doc_text}"""

    try:
        raw = call_groq(system, question, temperature=0.4)
        answer = clean_llm_output(raw)
        return jsonify({"answer": answer})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ────────────────────────────────────────────────────────────────────────────
# API – Stage 3 : Project Q&A Chatbot
# ────────────────────────────────────────────────────────────────────────────

PROJECT_SYSTEM = """You are an expert AI assistant for the NeuralMind AI project — a full-stack web application built with Flask and Groq's ultra-fast LLM inference.

## Project Overview
NeuralMind AI is a smart AI platform with three distinct AI-powered stages:

**Stage 1 — Emotion Detection Engine**
- Accepts any free-form text input
- Uses Qwen3-32B via Groq to identify 7 emotions: joy, sadness, anger, fear, surprise, disgust, neutral
- Returns confidence scores for each emotion, overall sentiment (positive/negative/neutral), intensity level, and an empathetic one-line insight
- Uses structured JSON output prompting with zero-shot classification

**Stage 2 — RAG Document Intelligence**
- Accepts PDF, DOCX, and TXT file uploads (up to 10MB)
- Extracts text using PyPDF2 (for PDFs) and python-docx (for Word documents)
- Stores extracted text per-session (up to 15,000 characters)
- Answers user questions strictly grounded in the uploaded document — no hallucination
- Uses a context-injection RAG (Retrieval-Augmented Generation) pattern

**Stage 3 — Project Q&A Chatbot (this stage)**
- Multi-turn conversational AI that knows the NeuralMind project deeply
- Maintains rolling 10-turn conversation history per session
- Scoped to project-related knowledge

## Technology Stack
- **Backend**: Python 3.10+, Flask web framework
- **AI Inference**: Groq API (ultra-fast LPU-based inference)
- **Model**: qwen/qwen3-32b (32 billion parameter model)
- **Document Parsing**: PyPDF2 (PDF), python-docx (DOCX), built-in (TXT)
- **Config Management**: python-dotenv
- **Frontend**: Vanilla HTML5, CSS3, JavaScript (no frameworks)
- **Session Management**: Flask sessions with in-memory stores

## Configuration
The `.env` file uses an inline comment convention:
```
GROQ_API_KEY_QWEN=gsk_XXXXXXXXX#qwen/qwen3-32b_emotion_llm
```
The app splits on `#` — left side is the API key, right side is the model ID.

## Architecture
Browser → Flask Route Handlers → Groq API (Qwen3-32B) → JSON Response → Frontend Rendering

## Deployment
- Local: `python app.py` (runs on port 5000)
- Production: Any WSGI-compatible host — Render, Railway, HuggingFace Spaces, Heroku, VPS with Gunicorn

## API Endpoints
- `POST /api/emotion` — emotion analysis
- `POST /api/upload` — document upload and text extraction
- `POST /api/rag` — document-grounded Q&A
- `POST /api/chat` — project chatbot

Answer questions clearly, helpfully, and in a conversational tone. Use markdown-style formatting in responses when it helps clarity (bullet points, bold terms, code snippets). If asked about something outside the project scope, politely acknowledge it and offer to help with project-related questions instead."""


@app.route("/api/chat", methods=["POST"])
def api_chat():
    data = request.get_json(silent=True) or {}
    message = (data.get("message") or "").strip()
    sid = session.get("sid") or os.urandom(8).hex()
    session["sid"] = sid

    if not message:
        return jsonify({"error": "No message provided"}), 400

    history = conv_store.setdefault(sid, [])
    history.append({"role": "user", "content": message})

    trimmed = history[-10:]

    try:
        response = client.chat.completions.create(
            model=model_name,
            messages=[{"role": "system", "content": PROJECT_SYSTEM}] + trimmed,
            temperature=0.7,
            max_tokens=768,
        )
        reply = response.choices[0].message.content.strip()
        # Clean think tags from chat too
        reply = clean_llm_output(reply)
        history.append({"role": "assistant", "content": reply})
        return jsonify({"reply": reply})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    app.run(debug=True, port=5000)